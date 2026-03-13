from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading(doc, text, level=1, color=(0, 51, 102)):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(*color)
    return heading

def add_table_row(table, col1, col2, bold_col1=True):
    row = table.add_row()
    row.cells[0].text = col1
    row.cells[1].text = col2
    if bold_col1:
        row.cells[0].paragraphs[0].runs[0].bold = True

def set_table_style(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

# Map area names to page ranges in inspection PDF
AREA_IMAGE_MAP = {
    "hall": range(27, 28),
    "bedroom": range(28, 30),
    "passage": range(29, 30),
    "staircase": range(30, 31),
    "master bedroom": range(31, 33),
    "external wall": range(34, 36),
    "common bathroom": range(34, 35),
    "balcony": range(35, 36),
    "terrace": range(36, 37),
}

def insert_images_for_area(doc, area_name, image_dir, max_images=2):
    area_lower = area_name.lower()
    matched_pages = []
    
    for key, pages in AREA_IMAGE_MAP.items():
        if key in area_lower:
            matched_pages = list(pages)
            break
    
    inserted = 0
    all_images = sorted(os.listdir(image_dir))
    
    for img_file in all_images:
        if inserted >= max_images:
            break
        # Extract page number from filename
        import re
        match = re.search(r'page(\d+)', img_file)
        if match:
            page_num = int(match.group(1))
            if page_num in matched_pages and 'inspection' in img_file:
                img_path = os.path.join(image_dir, img_file)
                try:
                    doc.add_picture(img_path, width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    inserted += 1
                except Exception:
                    pass
    
    if inserted == 0:
        doc.add_paragraph("Image Not Available").italic = True
def build_ddr(ddr_json, image_dir, output_path):
    doc = Document()

    # Title
    title = doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Prepared by: AI Report Generation System")
    doc.add_paragraph("─" * 60)

    # ── Section 1: Property Issue Summary ──
    add_heading(doc, "1. Property Issue Summary")
    doc.add_paragraph(ddr_json.get("property_issue_summary", "Not Available"))
    doc.add_paragraph()

    # ── Section 2: Area-wise Observations ──
    add_heading(doc, "2. Area-wise Observations")
    observations = ddr_json.get("area_wise_observations", [])

    for obs in observations:
        area = obs.get("area", "Unknown Area")
        observation = obs.get("observation", "Not Available")
        source = obs.get("source", "")

        add_heading(doc, area, level=2, color=(0, 102, 204))
        doc.add_paragraph(f"Observation: {observation}")
        doc.add_paragraph(f"Source: {source.title()}")

        # Embed relevant images
        insert_images_for_area(doc, area, image_dir, max_images=2)
        doc.add_paragraph()

    # ── Section 3: Probable Root Cause ──
    add_heading(doc, "3. Probable Root Cause")
    causes = ddr_json.get("probable_root_cause", [])
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Issue"
    table.rows[0].cells[1].text = "Root Cause"
    for row in table.rows[0].cells:
        for run in row.paragraphs[0].runs:
            run.bold = True
    for item in causes:
        add_table_row(table, item.get("issue", ""), item.get("cause", "Not Available"))
    set_table_style(table)
    doc.add_paragraph()

    # ── Section 4: Severity Assessment ──
    add_heading(doc, "4. Severity Assessment")
    severity_list = ddr_json.get("severity_assessment", [])
    table2 = doc.add_table(rows=1, cols=3)
    table2.rows[0].cells[0].text = "Area"
    table2.rows[0].cells[1].text = "Severity"
    table2.rows[0].cells[2].text = "Reasoning"
    for cell in table2.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for item in severity_list:
        row = table2.add_row()
        row.cells[0].text = item.get("area", "")
        row.cells[1].text = item.get("severity", "")
        row.cells[2].text = item.get("reasoning", "Not Available")
    set_table_style(table2)
    doc.add_paragraph()

    # ── Section 5: Recommended Actions ──
    add_heading(doc, "5. Recommended Actions")
    actions = ddr_json.get("recommended_actions", [])
    for item in actions:
        area = item.get("area", "")
        action = item.get("action", "Not Available")
        doc.add_paragraph(f"• {area}: {action}")
    doc.add_paragraph()

    # ── Section 6: Additional Notes ──
    add_heading(doc, "6. Additional Notes")
    doc.add_paragraph(ddr_json.get("additional_notes", "Not Available"))
    doc.add_paragraph()

    # ── Section 7: Missing / Unclear Information ──
    add_heading(doc, "7. Missing or Unclear Information")
    missing = ddr_json.get("missing_or_unclear_information", ["Not Available"])
    for item in missing:
        doc.add_paragraph(f"• {item}")

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"DDR saved to: {output_path}")